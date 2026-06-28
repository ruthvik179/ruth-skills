from __future__ import annotations

from html import escape
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.shared import Pt
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def render_resume_package(selected_path: str | Path, output_dir: str | Path, max_pages: int) -> dict[str, Any]:
    selected = yaml.safe_load(Path(selected_path).read_text(encoding="utf-8"))
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    docx_path = output / "resume.docx"
    pdf_path = output / "resume.pdf"
    _write_docx(selected, docx_path)
    _write_pdf(selected, pdf_path)

    page_count = count_pdf_pages(pdf_path)
    pdf_text = extract_pdf_text(pdf_path)
    expected_terms = [
        selected.get("job", {}).get("title", ""),
        *[company.get("company", "") for company in selected.get("experience", [])],
    ]
    pdf_text_check = "pass" if all(term in pdf_text for term in expected_terms if term) else "fail"
    metadata = {
        "docx": str(docx_path.name),
        "pdf": str(pdf_path.name),
        "page_count": page_count,
        "max_pages": max_pages,
        "page_count_check": "pass" if page_count <= max_pages else "fail",
        "pdf_text_check": pdf_text_check,
    }
    (output / "render.yaml").write_text(yaml.safe_dump(metadata, sort_keys=False), encoding="utf-8")
    if page_count > max_pages:
        raise ValueError(f"rendered PDF has {page_count} pages, max is {max_pages}")
    if pdf_text_check != "pass":
        raise ValueError("PDF text extraction did not contain expected resume terms")
    return metadata


def count_pdf_pages(path: str | Path) -> int:
    return len(PdfReader(str(path)).pages)


def extract_pdf_text(path: str | Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _write_docx(selected: dict[str, Any], path: Path) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    profile_title = selected.get("job", {}).get("title") or "Tailored Resume"
    document.add_heading(profile_title, level=0)
    _add_section(document, "Summary", [selected.get("summary", {}).get("text", "")])
    _add_section(document, "Skills", _skill_lines(selected))
    _add_experience(document, selected)
    _add_section(document, "Education", [item.get("text", "") for item in selected.get("education", [])])
    _add_section(document, "Projects", [item.get("text", "") for item in selected.get("projects", [])])
    _add_section(document, "Certifications", [item.get("text", "") for item in selected.get("certifications", [])])
    _add_section(document, "Extracurriculars", [item.get("text", "") for item in selected.get("extracurriculars", [])])
    document.save(path)


def _write_pdf(selected: dict[str, Any], path: Path) -> None:
    styles = getSampleStyleSheet()
    story = []
    title = selected.get("job", {}).get("title") or "Tailored Resume"
    story.append(Paragraph(_pdf_text(title), styles["Title"]))
    story.extend(_pdf_section("Summary", [selected.get("summary", {}).get("text", "")], styles))
    story.extend(_pdf_section("Skills", _skill_lines(selected), styles))
    story.extend(_pdf_experience(selected, styles))
    story.extend(_pdf_section("Education", [item.get("text", "") for item in selected.get("education", [])], styles))
    story.extend(_pdf_section("Projects", [item.get("text", "") for item in selected.get("projects", [])], styles))
    story.extend(_pdf_section("Certifications", [item.get("text", "") for item in selected.get("certifications", [])], styles))
    story.extend(_pdf_section("Extracurriculars", [item.get("text", "") for item in selected.get("extracurriculars", [])], styles))
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    doc.build(story)


def _skill_lines(selected: dict[str, Any]) -> list[str]:
    lines = []
    for group in selected.get("skills", {}).get("groups", []):
        lines.append(f"{group.get('name', '')}: {', '.join(group.get('items', []))}")
    return lines


def _add_section(document: Document, title: str, lines: list[str]) -> None:
    clean_lines = [line for line in lines if line]
    if not clean_lines:
        return
    document.add_heading(title, level=1)
    for line in clean_lines:
        document.add_paragraph(line)


def _add_experience(document: Document, selected: dict[str, Any]) -> None:
    if not selected.get("experience"):
        return
    document.add_heading("Experience", level=1)
    for company in selected.get("experience", []):
        document.add_heading(company.get("company", ""), level=2)
        for role in company.get("roles", []):
            document.add_paragraph(f"{role.get('title', '')} | {role.get('start', '')} - {role.get('end', '')}")
            for bullet in role.get("bullets", []):
                document.add_paragraph(bullet.get("text", ""), style="List Bullet")


def _pdf_section(title: str, lines: list[str], styles: Any) -> list[Any]:
    clean_lines = [line for line in lines if line]
    if not clean_lines:
        return []
    story: list[Any] = [Paragraph(_pdf_text(title), styles["Heading2"])]
    for line in clean_lines:
        story.append(Paragraph(_pdf_text(line), styles["BodyText"]))
    story.append(Spacer(1, 8))
    return story


def _pdf_experience(selected: dict[str, Any], styles: Any) -> list[Any]:
    if not selected.get("experience"):
        return []
    story: list[Any] = [Paragraph("Experience", styles["Heading2"])]
    for company in selected.get("experience", []):
        story.append(Paragraph(_pdf_text(company.get("company", "")), styles["Heading3"]))
        for role in company.get("roles", []):
            story.append(
                Paragraph(
                    _pdf_text(f"{role.get('title', '')} | {role.get('start', '')} - {role.get('end', '')}"),
                    styles["BodyText"],
                )
            )
            for bullet in role.get("bullets", []):
                story.append(Paragraph(_pdf_text(f"- {bullet.get('text', '')}"), styles["BodyText"]))
    story.append(Spacer(1, 8))
    return story


def _pdf_text(value: Any) -> str:
    return escape(str(value), quote=False)
